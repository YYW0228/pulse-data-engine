"""
scripts/doc_parser.py — 数据接入包解析器 (方法D: 数据适配流水线)

输入: 任意 PDF/Word/扫描件/网页/md/txt
输出: 标准化块 (可喂给 compliance_index.py 的块格式)

这是"客户私有文档 → 可索引、可引用、可防护知识"的物理前提。
每个新客户/新场景的数据替换成本从"50天人工"→"流水线+少量人工校验"。

用法:
  uv run python -m scripts.doc_parser <文件或目录> [--format auto|pdf|docx|md|txt]

依赖: pymupdf (PDF), python-docx (Word), 均轻量。
"""

from __future__ import annotations

import argparse
import hashlib
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))


def parse_file(path: Path) -> str:
    """按扩展名解析文档 → 纯文本"""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in (".docx", ".doc"):
        return _parse_docx(path)
    if ext in (".md", ".txt", ".text"):
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def _parse_pdf(path: Path) -> str:
    """PDF 解析 (pymupdf — 支持扫描件 OCR 前的基础文本层)"""
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n\n".join(parts)


def _parse_docx(path: Path) -> str:
    """Word 解析 (python-docx)"""
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    # 表格
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def split_into_chunks(text: str, min_chars: int = 200) -> list[dict]:
    """按标题结构分块 (复用 compliance_index 的块格式)

    规则: ## / ### 标题切分; 无标题时按段落聚合
    返回: [{doc_name, title, content, char_len, importance}]
    """
    lines = text.splitlines()
    blocks: list[dict] = []
    current_title = "文档头部"
    current: list[str] = []

    def flush() -> None:
        if current:
            content = "\n".join(current).strip()
            if len(content) >= min_chars:
                blocks.append({
                    "doc_name": "unknown",  # 调用方填
                    "title": current_title,
                    "content": content,
                    "char_len": len(content),
                    "importance": _importance(current_title, content),
                })
            elif current:  # 短块也保留 (重要信息常在短段)
                blocks.append({
                    "doc_name": "unknown",
                    "title": current_title,
                    "content": content,
                    "char_len": len(content),
                    "importance": _importance(current_title, content) * 0.8,
                })
            current.clear()

    for line in lines:
        stripped = line.strip()
        if re.match(r"^#{1,6}\s", stripped):  # 标题
            flush()
            current_title = re.sub(r"^#+\s*", "", stripped)
        else:
            current.append(line)
    flush()
    return blocks


def _importance(title: str, content: str) -> float:
    """重要性启发式 (与 compliance_index 对齐)"""
    score = 0.5
    if re.match(r"^#\s", title) or title in ("文档头部",):
        score = 0.6
    if re.match(r"^#{2}\s", title):
        score = 0.8
    if re.match(r"^#{3,}\s", title):
        score = 0.7
    # 实体密度: 条款/金额/日期 → 高价值
    entity_hits = len(re.findall(r"第.+条|第.+款|[0-9,]+(?:万|元|%)|20\d{2}年", content))
    if entity_hits >= 3:
        score = min(score + 0.1, 1.0)
    return round(score, 2)


def main():
    parser = argparse.ArgumentParser(description="数据接入包文档解析器")
    parser.add_argument("path", help="文件或目录")
    parser.add_argument("--out", default="data/parsed", help="输出 JSON 目录")
    args = parser.parse_args()

    src = Path(args.path)
    files = [src] if src.is_file() else sorted(src.rglob("*")) if src.is_dir() else []
    files = [f for f in files if f.suffix.lower() in (".pdf", ".docx", ".doc", ".md", ".txt")]

    if not files:
        print(f"未找到可解析文件: {args.path}")
        sys.exit(1)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    total_blocks = 0
    total_chars = 0

    for f in files:
        try:
            text = parse_file(f)
            if not text.strip():
                print(f"  ⚠️ {f.name}: 无文本层 (可能需 OCR)")
                continue
            blocks = split_into_chunks(text)
            for b in blocks:
                b["doc_name"] = f.name
                b["doc_hash"] = hashlib.sha256(b["content"].encode()).hexdigest()[:16]
            # 输出 JSONL
            out = out_dir / (f.stem + ".jsonl")
            import json

            with out.open("w", encoding="utf-8") as fh:
                for b in blocks:
                    fh.write(json.dumps(b, ensure_ascii=False) + "\n")
            total_blocks += len(blocks)
            total_chars += sum(b["char_len"] for b in blocks)
            print(f"  ✅ {f.name}: {len(blocks)} 块, {sum(b['char_len'] for b in blocks)} 字符 → {out.name}")
        except Exception as e:
            print(f"  ❌ {f.name}: {e}")

    print(f"\n总计: {len(files)} 文件, {total_blocks} 块, {total_chars} 字符")
    print(f"输出: {out_dir}/ (JSONL, 可直接喂 compliance_index)")


if __name__ == "__main__":
    main()
